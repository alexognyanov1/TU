import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "solution.docx")

A1, A2 = 3.0, 2.0
a1, a2 = 1.0, 0.5


def generate_plots():
    tau = np.linspace(-6, 6, 1000)

    R11 = (A1**2 / (2 * a1)) * np.exp(-a1 * np.abs(tau))
    R22 = (A2**2 / (2 * a2)) * np.exp(-a2 * np.abs(tau))
    R21 = np.where(
        tau >= 0,
        (A1 * A2 / (a1 + a2)) * np.exp(-a2 * tau),
        (A1 * A2 / (a1 + a2)) * np.exp(a1 * tau),
    )

    fig, axes = plt.subplots(3, 1, figsize=(8, 10))

    axes[0].plot(tau, R11, 'b-', linewidth=2)
    axes[0].set_title(r'$R_{11}(\tau) = \frac{A_1^2}{2a_1} e^{-a_1|\tau|}$', fontsize=14)
    axes[0].set_xlabel(r'$\tau$', fontsize=12)
    axes[0].set_ylabel(r'$R_{11}(\tau)$', fontsize=12)
    axes[0].grid(True, alpha=0.3)
    axes[0].axhline(y=0, color='k', linewidth=0.5)
    axes[0].axvline(x=0, color='k', linewidth=0.5)
    r11_0 = A1**2 / (2 * a1)
    axes[0].annotate(f'R₁₁(0) = {r11_0:.1f}', xy=(0, r11_0),
                     xytext=(1.5, r11_0 * 0.8), fontsize=10,
                     arrowprops=dict(arrowstyle='->', color='red'))

    axes[1].plot(tau, R22, 'r-', linewidth=2)
    axes[1].set_title(r'$R_{22}(\tau) = \frac{A_2^2}{2a_2} e^{-a_2|\tau|}$', fontsize=14)
    axes[1].set_xlabel(r'$\tau$', fontsize=12)
    axes[1].set_ylabel(r'$R_{22}(\tau)$', fontsize=12)
    axes[1].grid(True, alpha=0.3)
    axes[1].axhline(y=0, color='k', linewidth=0.5)
    axes[1].axvline(x=0, color='k', linewidth=0.5)
    r22_0 = A2**2 / (2 * a2)
    axes[1].annotate(f'R₂₂(0) = {r22_0:.1f}', xy=(0, r22_0),
                     xytext=(1.5, r22_0 * 0.8), fontsize=10,
                     arrowprops=dict(arrowstyle='->', color='red'))

    axes[2].plot(tau, R21, 'g-', linewidth=2)
    axes[2].set_title(r'$R_{21}(\tau)$ — Cross-correlation', fontsize=14)
    axes[2].set_xlabel(r'$\tau$', fontsize=12)
    axes[2].set_ylabel(r'$R_{21}(\tau)$', fontsize=12)
    axes[2].grid(True, alpha=0.3)
    axes[2].axhline(y=0, color='k', linewidth=0.5)
    axes[2].axvline(x=0, color='k', linewidth=0.5)
    r21_0 = A1 * A2 / (a1 + a2)
    axes[2].annotate(f'R₂₁(0) = {r21_0:.1f}', xy=(0, r21_0),
                     xytext=(1.5, r21_0 * 0.8), fontsize=10,
                     arrowprops=dict(arrowstyle='->', color='red'))

    plt.tight_layout()
    plot_path = os.path.join(OUTPUT_DIR, "correlation_plots.png")
    plt.savefig(plot_path, dpi=200, bbox_inches='tight')
    plt.close()

    fig2, axes2 = plt.subplots(1, 2, figsize=(10, 4))
    t = np.linspace(0, 6, 500)
    s1 = A1 * np.exp(-a1 * t)
    s2 = A2 * np.exp(-a2 * t)

    axes2[0].plot(t, s1, 'b-', linewidth=2)
    axes2[0].set_title(r'$S_1(t) = A_1 \cdot e^{-a_1 t}$', fontsize=13)
    axes2[0].set_xlabel('t', fontsize=12)
    axes2[0].set_ylabel(r'$S_1(t)$', fontsize=12)
    axes2[0].grid(True, alpha=0.3)
    axes2[0].set_xlim(-0.5, 6)
    axes2[0].set_ylim(-0.2, A1 + 0.5)
    axes2[0].axhline(y=0, color='k', linewidth=0.5)
    axes2[0].axvline(x=0, color='k', linewidth=0.5)

    axes2[1].plot(t, s2, 'r-', linewidth=2)
    axes2[1].set_title(r'$S_2(t) = A_2 \cdot e^{-a_2 t}$', fontsize=13)
    axes2[1].set_xlabel('t', fontsize=12)
    axes2[1].set_ylabel(r'$S_2(t)$', fontsize=12)
    axes2[1].grid(True, alpha=0.3)
    axes2[1].set_xlim(-0.5, 6)
    axes2[1].set_ylim(-0.2, A2 + 0.5)
    axes2[1].axhline(y=0, color='k', linewidth=0.5)
    axes2[1].axvline(x=0, color='k', linewidth=0.5)

    plt.tight_layout()
    signals_path = os.path.join(OUTPUT_DIR, "signals.png")
    plt.savefig(signals_path, dpi=200, bbox_inches='tight')
    plt.close()

    return plot_path, signals_path


def add_paragraph(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=Pt(6), space_before=Pt(0), font_name='Times New Roman'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    return p


def _math_run(text, bold=False, size=None):
    r = OxmlElement('m:r')
    wrpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Cambria Math')
    rfonts.set(qn('w:hAnsi'), 'Cambria Math')
    wrpr.append(rfonts)
    if bold:
        wrpr.append(OxmlElement('w:b'))
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))
        wrpr.append(sz)
    r.append(wrpr)
    t = OxmlElement('m:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _math_nary(sub_text, sup_text, body_elements=None, op='∫', bold=False, size=None):
    nary = OxmlElement('m:nary')
    naryPr = OxmlElement('m:naryPr')
    chr_el = OxmlElement('m:chr')
    chr_el.set(qn('m:val'), op)
    naryPr.append(chr_el)
    limLoc = OxmlElement('m:limLoc')
    limLoc.set(qn('m:val'), 'subSup')
    naryPr.append(limLoc)
    nary.append(naryPr)
    sub_el = OxmlElement('m:sub')
    sub_el.append(_math_run(sub_text, bold=bold, size=size))
    nary.append(sub_el)
    sup_el = OxmlElement('m:sup')
    sup_el.append(_math_run(sup_text, bold=bold, size=size))
    nary.append(sup_el)
    e_el = OxmlElement('m:e')
    if body_elements:
        for be in body_elements:
            e_el.append(be)
    nary.append(e_el)
    return nary


def _math_sSup(base_text, sup_text, bold=False, size=None):
    sSup = OxmlElement('m:sSup')
    e_el = OxmlElement('m:e')
    e_el.append(_math_run(base_text, bold=bold, size=size))
    sSup.append(e_el)
    sup_el = OxmlElement('m:sup')
    sup_el.append(_math_run(sup_text, bold=bold, size=size))
    sSup.append(sup_el)
    return sSup


_UNI_SUB = {'₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5',
            '₆': '6', '₇': '7', '₈': '8', '₉': '9', '₊': '+', '₋': '-'}


def _parse_math(s, bold=False, size=None):
    elements = []
    i = 0
    n = len(s)
    text_buf = ['']

    def flush():
        if text_buf[0]:
            elements.append(_math_run(text_buf[0], bold=bold, size=size))
            text_buf[0] = ''

    def read_paren_group(s, j):
        depth, j = 1, j + 1
        out = ''
        while j < n and depth > 0:
            c = s[j]
            if c == '(':
                depth += 1
                out += c
            elif c == ')':
                depth -= 1
                if depth > 0:
                    out += c
            else:
                out += c
            j += 1
        return out, j

    while i < n:
        ch = s[i]

        if ch == '∫':
            flush()
            j = i + 1
            sub = ''
            if j < n and s[j] == '_':
                j += 1
                if j < n and s[j] == '(':
                    sub, j = read_paren_group(s, j)
                else:
                    while j < n and s[j] != '^':
                        sub += s[j]
                        j += 1
            else:
                while j < n and s[j] != '^':
                    c = s[j]
                    if c in _UNI_SUB:
                        sub += _UNI_SUB[c]
                        j += 1
                    elif c == '∞' and sub:
                        sub += '∞'
                        j += 1
                    else:
                        break

            sup = ''
            if j < n and s[j] == '^':
                j += 1
                if j < n and s[j] == '(':
                    sup, j = read_paren_group(s, j)
                else:
                    while j < n and s[j] not in ' \t,':
                        sup += s[j]
                        j += 1

            body_elements = _parse_math(s[j:], bold=bold, size=size)
            elements.append(_math_nary(sub, sup, body_elements=body_elements,
                                       bold=bold, size=size))
            return elements

        if ch == 'e' and i + 1 < n and s[i + 1] == '^':
            flush()
            j = i + 2
            sup = ''
            if j < n and s[j] == '(':
                sup, j = read_paren_group(s, j)
            else:
                while j < n and s[j] not in ' \t,)':
                    sup += s[j]
                    j += 1
            elements.append(_math_sSup('e', sup, bold=bold, size=size))
            i = j
            continue

        text_buf[0] += ch
        i += 1

    flush()
    return elements


def add_formula(doc, text, size=12, bold=False, space_before=Pt(6), space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before

    elements = _parse_math(text, bold=bold, size=size)

    oMathPara = OxmlElement('m:oMathPara')
    oMathParaPr = OxmlElement('m:oMathParaPr')
    jc = OxmlElement('m:jc')
    jc.set(qn('m:val'), 'center')
    oMathParaPr.append(jc)
    oMathPara.append(oMathParaPr)
    oMath = OxmlElement('m:oMath')
    for el in elements:
        oMath.append(el)
    oMathPara.append(oMath)
    p._p.append(oMathPara)
    return p


def build_document(plot_path, signals_path):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === TITLE PAGE ===
    for _ in range(4):
        add_paragraph(doc, '', size=12)

    add_paragraph(doc, 'ТЕХНИЧЕСКИ УНИВЕРСИТЕТ', bold=True, size=16,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_paragraph(doc, 'Факултет по електроника и автоматика', size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    add_paragraph(doc, 'КУРСОВА РАБОТА', bold=True, size=18,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_paragraph(doc, 'по', size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_paragraph(doc, 'Сигнали и Системи', bold=True, size=16,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

    add_paragraph(doc, 'Тема: Корелационни функции на аналогови непериодични сигнали',
                  bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

    add_paragraph(doc, 'Изготвил: ......................................', size=12,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(8))
    add_paragraph(doc, 'Специалност: ..................................', size=12,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(8))
    add_paragraph(doc, 'Фак. №: ..........................................', size=12,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(8))
    add_paragraph(doc, 'Група: ............................................', size=12,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(8))

    for _ in range(4):
        add_paragraph(doc, '', size=12)

    add_paragraph(doc, 'Оценка: .................    Проверил: .................', size=12,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    doc.add_page_break()

    # === PAGE 2: ASSIGNMENT ===
    add_paragraph(doc, 'Задание', bold=True, size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_paragraph(doc, 'Дадени са следните аналогови непериодични сигнали:', size=12, space_after=Pt(6))

    add_formula(doc, 'S₁(t) = A₁ · exp(-a₁ · t),    t ≥ 0')
    add_formula(doc, 'S₂(t) = A₂ · exp(-a₂ · t),    t ≥ 0')

    doc.add_picture(signals_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, '', size=6)
    add_paragraph(doc, 'а) Да се определи изразът за автокорелационната функция на първия сигнал.', size=12, space_after=Pt(4))
    add_paragraph(doc, 'б) Да се определи изразът за автокорелационната функция на втория сигнал.', size=12, space_after=Pt(4))
    add_paragraph(doc, 'в) Да се определи изразът за взаимнокорелационната функция на двата сигнала, ако сигналът който се измества във времето е първият сигнал.', size=12, space_after=Pt(4))
    add_paragraph(doc, 'г) Да се начертаят графично получените корелационни функции.', size=12, space_after=Pt(12))

    doc.add_page_break()

    # === PAGE 3+: SOLUTION ===
    add_paragraph(doc, 'Решение', bold=True, size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # --- Definitions ---
    add_paragraph(doc, '1. Използвани математически зависимости', bold=True, size=13, space_after=Pt(8), space_before=Pt(12))

    add_paragraph(doc, 'Автокорелационна функция на енергиен сигнал s(t):', size=12, space_after=Pt(4))
    add_formula(doc, 'Rss(τ) = ∫₋∞^∞  s(t) · s(t − τ) dt')

    add_paragraph(doc, 'Взаимнокорелационна функция на два енергийни сигнала:', size=12, space_after=Pt(4))
    add_formula(doc, 'R₂₁(τ) = ∫₋∞^∞  s₂(t) · s₁(t − τ) dt')

    add_paragraph(doc, 'където τ е времевото изместване, а s₁ е сигналът, който се измества.', size=12, space_after=Pt(4))

    add_paragraph(doc, 'Основен интеграл, използван в решението:', size=12, space_after=Pt(4))
    add_formula(doc, '∫₀^∞  e^(-αt) dt = 1/α,    α > 0')
    add_formula(doc, '∫_a^∞  e^(-αt) dt = e^(-αa) / α,    α > 0')

    # --- Task a) ---
    add_paragraph(doc, '2. Задача а) — Автокорелационна функция на S₁(t)', bold=True, size=13, space_after=Pt(8), space_before=Pt(16))

    add_paragraph(doc, 'Даден е сигналът:', size=12, space_after=Pt(4))
    add_formula(doc, 'S₁(t) = A₁ · e^(-a₁t),    t ≥ 0    (иначе S₁(t) = 0)')

    add_paragraph(doc, 'Прилагаме дефиницията за автокорелационна функция:', size=12, space_after=Pt(4))
    add_formula(doc, 'R₁₁(τ) = ∫₋∞^∞  S₁(t) · S₁(t − τ) dt')

    add_paragraph(doc, 'Интегрантът е ненулев само когато и двата аргумента са в областта на определение:', size=12, space_after=Pt(4))
    add_paragraph(doc, '• S₁(t) ≠ 0  ⟹  t ≥ 0', size=12, space_after=Pt(2))
    add_paragraph(doc, '• S₁(t − τ) ≠ 0  ⟹  t − τ ≥ 0  ⟹  t ≥ τ', size=12, space_after=Pt(8))

    # Case tau >= 0
    add_paragraph(doc, 'Случай 1: τ ≥ 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Когато τ ≥ 0, долната граница е max(0, τ) = τ:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₁₁(τ) = ∫_τ^∞  A₁·e^(-a₁t) · A₁·e^(-a₁(t−τ)) dt')
    add_formula(doc, '= A₁² · e^(a₁τ) · ∫_τ^∞  e^(-2a₁t) dt')
    add_formula(doc, '= A₁² · e^(a₁τ) · e^(-2a₁τ) / (2a₁)')
    add_formula(doc, '= A₁² / (2a₁) · e^(-a₁τ)')

    # Case tau < 0
    add_paragraph(doc, 'Случай 2: τ < 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Когато τ < 0, условието t ≥ τ е автоматично изпълнено за t ≥ 0, така че долната граница е 0:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₁₁(τ) = ∫₀^∞  A₁·e^(-a₁t) · A₁·e^(-a₁(t−τ)) dt')
    add_formula(doc, '= A₁² · e^(a₁τ) · ∫₀^∞  e^(-2a₁t) dt')
    add_formula(doc, '= A₁² · e^(a₁τ) · 1/(2a₁)')
    add_formula(doc, '= A₁² / (2a₁) · e^(a₁τ)')

    add_paragraph(doc, 'Тъй като за τ ≥ 0 имаме e^(-a₁τ), а за τ < 0 имаме e^(a₁τ), и двата случая се обединяват в:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₁₁(τ) = A₁² / (2a₁) · e^(-a₁|τ|)',
                size=13, bold=True,
                space_before=Pt(8), space_after=Pt(12))

    # --- Task b) ---
    add_paragraph(doc, '3. Задача б) — Автокорелационна функция на S₂(t)', bold=True, size=13, space_after=Pt(8), space_before=Pt(16))

    add_paragraph(doc, 'Даден е сигналът:', size=12, space_after=Pt(4))
    add_formula(doc, 'S₂(t) = A₂ · e^(-a₂t),    t ≥ 0    (иначе S₂(t) = 0)')

    add_paragraph(doc, 'Процедурата е аналогична на задача а). Прилагаме дефиницията:', size=12, space_after=Pt(4))
    add_formula(doc, 'R₂₂(τ) = ∫₋∞^∞  S₂(t) · S₂(t − τ) dt')

    add_paragraph(doc, 'Случай 1: τ ≥ 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Долната граница на интегриране е max(0, τ) = τ:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₂(τ) = ∫_τ^∞  A₂·e^(-a₂t) · A₂·e^(-a₂(t−τ)) dt')
    add_formula(doc, '= A₂² · e^(a₂τ) · ∫_τ^∞  e^(-2a₂t) dt')
    add_formula(doc, '= A₂² · e^(a₂τ) · e^(-2a₂τ) / (2a₂)')
    add_formula(doc, '= A₂² / (2a₂) · e^(-a₂τ)')

    add_paragraph(doc, 'Случай 2: τ < 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Условието t ≥ τ е автоматично изпълнено за t ≥ 0, така че долната граница е 0:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₂(τ) = ∫₀^∞  A₂·e^(-a₂t) · A₂·e^(-a₂(t−τ)) dt')
    add_formula(doc, '= A₂² · e^(a₂τ) · ∫₀^∞  e^(-2a₂t) dt')
    add_formula(doc, '= A₂² · e^(a₂τ) · 1/(2a₂)')
    add_formula(doc, '= A₂² / (2a₂) · e^(a₂τ)')

    add_paragraph(doc, 'Обединявайки двата случая:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₂(τ) = A₂² / (2a₂) · e^(-a₂|τ|)',
                size=13, bold=True,
                space_before=Pt(8), space_after=Pt(12))

    # --- Task c) ---
    add_paragraph(doc, '4. Задача в) — Взаимнокорелационна функция R₂₁(τ)', bold=True, size=13, space_after=Pt(8), space_before=Pt(16))

    add_paragraph(doc, 'Търсим взаимнокорелационната функция на S₂(t) и S₁(t), където S₁ е сигналът, който се измества във времето:', size=12, space_after=Pt(4))
    add_formula(doc, 'R₂₁(τ) = ∫₋∞^∞  S₂(t) · S₁(t − τ) dt')

    add_paragraph(doc, 'Условия за ненулев интегрант:', size=12, space_after=Pt(4))
    add_paragraph(doc, '• S₂(t) ≠ 0  ⟹  t ≥ 0', size=12, space_after=Pt(2))
    add_paragraph(doc, '• S₁(t − τ) ≠ 0  ⟹  t ≥ τ', size=12, space_after=Pt(8))

    add_paragraph(doc, 'Случай 1: τ ≥ 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Долната граница е max(0, τ) = τ:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₁(τ) = ∫_τ^∞  A₂·e^(-a₂t) · A₁·e^(-a₁(t−τ)) dt')
    add_formula(doc, '= A₁·A₂ · e^(a₁τ) · ∫_τ^∞  e^(-(a₁+a₂)t) dt')
    add_formula(doc, '= A₁·A₂ · e^(a₁τ) · e^(-(a₁+a₂)τ) / (a₁ + a₂)')
    add_formula(doc, '= A₁·A₂ / (a₁ + a₂) · e^(-a₂τ)')

    add_paragraph(doc, 'Случай 2: τ < 0', bold=True, size=12, space_after=Pt(4), space_before=Pt(8))
    add_paragraph(doc, 'Долната граница е max(0, τ) = 0 (тъй като τ < 0):', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₁(τ) = ∫₀^∞  A₂·e^(-a₂t) · A₁·e^(-a₁(t−τ)) dt')
    add_formula(doc, '= A₁·A₂ · e^(a₁τ) · ∫₀^∞  e^(-(a₁+a₂)t) dt')
    add_formula(doc, '= A₁·A₂ · e^(a₁τ) · 1/(a₁ + a₂)')
    add_formula(doc, '= A₁·A₂ / (a₁ + a₂) · e^(a₁τ)')

    add_paragraph(doc, 'Крайният резултат:', size=12, space_after=Pt(4))

    add_formula(doc, 'R₂₁(τ) = A₁·A₂ / (a₁+a₂) · e^(-a₂τ),    τ ≥ 0',
                size=13, bold=True,
                space_before=Pt(8), space_after=Pt(4))
    add_formula(doc, 'R₂₁(τ) = A₁·A₂ / (a₁+a₂) · e^(a₁τ),     τ < 0',
                size=13, bold=True,
                space_before=Pt(4), space_after=Pt(12))

    add_paragraph(doc, 'Забележка: За разлика от автокорелационните функции, взаимнокорелационната функция НЕ е четна. Тя е непрекъсната в τ = 0, където стойността е R₂₁(0) = A₁·A₂/(a₁+a₂), но скоростта на спадане е различна за положителни и отрицателни τ (определя се от a₂ и a₁ съответно).', size=11, italic=True, space_after=Pt(12))

    # --- Task d) Plots ---
    doc.add_page_break()
    add_paragraph(doc, '5. Задача г) — Графики на корелационните функции', bold=True, size=13, space_after=Pt(8), space_before=Pt(4))

    add_paragraph(doc, f'За графичното представяне са използвани следните числени стойности: A₁ = {A1}, A₂ = {A2}, a₁ = {a1}, a₂ = {a2}.', size=12, space_after=Pt(4))

    add_paragraph(doc, 'Стойности в τ = 0 (максимуми на автокорелационните функции):', size=12, space_after=Pt(4))
    add_formula(doc, f'R₁₁(0) = A₁²/(2a₁) = {A1}²/(2·{a1}) = {A1**2/(2*a1):.2f}')
    add_formula(doc, f'R₂₂(0) = A₂²/(2a₂) = {A2}²/(2·{a2}) = {A2**2/(2*a2):.2f}')
    add_formula(doc, f'R₂₁(0) = A₁·A₂/(a₁+a₂) = {A1}·{A2}/({a1}+{a2}) = {A1*A2/(a1+a2):.2f}')

    doc.add_picture(plot_path, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, 'Фиг. 1. Графики на автокорелационните функции R₁₁(τ), R₂₂(τ) и взаимнокорелационната функция R₂₁(τ).', size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_paragraph(doc, 'От графиките се наблюдава:', size=12, space_after=Pt(4))
    add_paragraph(doc, '• Автокорелационните функции R₁₁(τ) и R₂₂(τ) са четни (симетрични спрямо τ = 0) и имат максимум при τ = 0, което е характерно свойство на автокорелационните функции.', size=12, space_after=Pt(2))
    add_paragraph(doc, '• Взаимнокорелационната функция R₂₁(τ) НЕ е четна — скоростта на спадане е различна за τ > 0 (определя се от a₂) и за τ < 0 (определя се от a₁).', size=12, space_after=Pt(2))
    add_paragraph(doc, '• Всички корелационни функции спадат експоненциално към нула за |τ| → ∞.', size=12, space_after=Pt(12))

    # --- Appendix: Source code ---
    doc.add_page_break()
    add_paragraph(doc, 'Приложение — Програмен код (Python)', bold=True, size=13,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_paragraph(doc, 'За генериране на графиките е използван следният Python код:', size=12, space_after=Pt(8))

    code_text = '''import numpy as np
import matplotlib.pyplot as plt

A1, A2 = 3.0, 2.0
a1, a2 = 1.0, 0.5
tau = np.linspace(-6, 6, 1000)

# Autocorrelation R11
R11 = (A1**2 / (2*a1)) * np.exp(-a1 * np.abs(tau))

# Autocorrelation R22
R22 = (A2**2 / (2*a2)) * np.exp(-a2 * np.abs(tau))

# Cross-correlation R21
R21 = np.where(
    tau >= 0,
    (A1*A2 / (a1+a2)) * np.exp(-a2 * tau),
    (A1*A2 / (a1+a2)) * np.exp(a1 * tau),
)

fig, axes = plt.subplots(3, 1, figsize=(8, 10))
axes[0].plot(tau, R11, 'b-', linewidth=2)
axes[0].set_title('R11(tau)')
axes[0].grid(True)

axes[1].plot(tau, R22, 'r-', linewidth=2)
axes[1].set_title('R22(tau)')
axes[1].grid(True)

axes[2].plot(tau, R21, 'g-', linewidth=2)
axes[2].set_title('R21(tau)')
axes[2].grid(True)

plt.tight_layout()
plt.savefig("correlation_plots.png", dpi=200)
plt.show()'''

    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    pf = p.paragraph_format
    pf.space_after = Pt(6)

    doc.save(DOCX_PATH)
    print(f"Document saved to: {DOCX_PATH}")


if __name__ == "__main__":
    plot_path, signals_path = generate_plots()
    build_document(plot_path, signals_path)
    print("Done!")
